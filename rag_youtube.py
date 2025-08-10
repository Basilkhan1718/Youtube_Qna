import os
from langchain_core.documents import Document
import webvtt
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnableParallel, RunnableLambda, RunnablePassthrough

# Uncomment and run once if you want to download subtitles programmatically
# import subprocess
# video_id = "qJeaCHQ1k2w"
# cmd = f"yt-dlp --write-auto-sub --sub-lang en --skip-download --output \"%(id)s.%(ext)s\" https://www.youtube.com/watch?v={video_id}"
# subprocess.run(cmd, shell=True, check=True)

def parse_vtt_to_chunks(vtt_filename="qJeaCHQ1k2w.en.vtt", max_chars=1000):
    captions = []
    for caption in webvtt.read(vtt_filename):
        captions.append({
            "text": caption.text,
            "start": caption.start,
            "end": caption.end,
        })

    grouped_chunks = []
    group = []
    char_count = 0

    for cap in captions:
        group.append(cap)
        char_count += len(cap["text"])
        if char_count >= max_chars:
            grouped_chunks.append(
                Document(
                    page_content=" ".join(c["text"] for c in group),
                    metadata={
                        "start": group[0]["start"],
                        "end": group[-1]["end"],
                    },
                )
            )
            group = []
            char_count = 0

    if group:
        grouped_chunks.append(
            Document(
                page_content=" ".join(c["text"] for c in group),
                metadata={
                    "start": group[0]["start"],
                    "end": group[-1]["end"],
                },
            )
        )

    transcript = " ".join([cap["text"] for cap in captions])
    return transcript, grouped_chunks

def build_pipeline(vtt_filename="qJeaCHQ1k2w.en.vtt", max_chars=1000):
    # Parse subtitles to documents with timestamps
    transcript, grouped_chunks = parse_vtt_to_chunks(vtt_filename, max_chars)

	 splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    fine_chunks = splitter.split_documents(grouped_chunks)

    # Create embeddings and vector store
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.from_documents(grouped_chunks, embeddings)
    retriever = vector_store.as_retriever(search_type="similarity", search_kwargs={"k": 4})

    # Init Gemini chat model
    llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.2)

    # Prompt & chain
    prompt = PromptTemplate(
        template="""
You are a helpful assistant.
Answer ONLY from the provided transcript context.
If the context is insufficient, just say you don't know.

Answer in {language}.

{context}
Question: {question}
""",
        input_variables=["context", "question", "language"]
    )

    def format_docs(docs):
        return "\n\n".join(doc.page_content for doc in docs)

    get_context_chain = (
        RunnableLambda(lambda x: retriever.get_relevant_documents(x["question"])) |
        RunnableLambda(format_docs)
    )

    main_chain = RunnableParallel({
        "context": get_context_chain,
        "question": RunnableLambda(lambda x: x["question"]),
        "language": RunnableLambda(lambda x: x["language"])
    }) | prompt | llm | StrOutputParser()

    return {
        "transcript": transcript,
        "grouped_chunks": grouped_chunks,
        "vector_store": vector_store,
        "retriever": retriever,
        "llm": llm,
        "main_chain": main_chain,
    }

def timestamp_link(start, video_id="qJeaCHQ1k2w"):
    parts = start.split(":")
    seconds = int(parts[0])*3600 + int(parts[1])*60 + int(float(parts[2]))
    url = f"https://www.youtube.com/watch?v={video_id}&t={seconds}s"
    label = f"{int(parts[0]):02}:{int(parts[1]):02}:{int(float(parts[2])):02}"
    return url, label

def search_transcript(query, retriever, video_id="qJeaCHQ1k2w"):
    results = retriever.get_relevant_documents(query)
    snippets = []
    for doc in results:
        start_time = doc.metadata.get("start", None)
        if start_time:
            ts_url, ts_label = timestamp_link(start_time, video_id)
            snippet = f"At [{ts_label}]({ts_url}): {doc.page_content[:180].replace(chr(10), ' ')}..."
        else:
            snippet = doc.page_content[:180].replace(chr(10), " ") + "..."
        snippets.append(snippet)
    return snippets

def generate_quiz_gemini(llm, text, n_questions=5):
    prompt = f"""
You are a creative quiz maker.
From the following context, create {n_questions} multiple choice questions with 4 options each and mark the correct answer.
Format:

Q1. Question?
A. Option 1
B. Option 2
C. Option 3
D. Option 4

Only output the quiz.

Context: {text}
"""
    return llm.invoke(prompt)

def export_text_to_pdf(text, filename="output.pdf"):
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in text.split("\n"):
        pdf.multi_cell(0, 10, line)
    pdf.output(filename)
    print(f"Saved PDF as {filename}")

def export_text_to_docx(text, filename="output.docx"):
    from docx import Document as DocxDoc

    doc = DocxDoc()
    doc.add_paragraph(text)
    doc.save(filename)
    print(f"Saved DOCX as {filename}")

def export_text_to_csv(chunks, filename="transcript_chunks.csv"):
    import csv
    with open(filename, "w", newline='', encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Start", "End", "Text"])
        for doc in chunks:
            writer.writerow([doc.metadata.get("start", ""), doc.metadata.get("end", ""), doc.page_content])
    print(f"Saved CSV as {filename}")

def text_to_speech(text, filename="output.mp3", lang="en"):
    from gtts import gTTS
    import IPython.display as ipd

    tts = gTTS(text=text, lang=lang)
    tts.save(filename)
    print(f"Audio saved as {filename}")
    return ipd.Audio(filename)
